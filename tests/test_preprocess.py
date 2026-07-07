from pathlib import Path

import pytest

import preprocess


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(0, 10, "Hello GraphRAG preprocessing smoke test.")
    dest = tmp_path / "sample.pdf"
    pdf.output(str(dest))
    return dest


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    import docx

    document = docx.Document()
    document.add_paragraph("안녕하세요 GraphRAG 전처리 테스트입니다.")
    dest = tmp_path / "sample.docx"
    document.save(str(dest))
    return dest


def test_clean_text_removes_control_chars_and_extra_whitespace():
    dirty = "Hello\x00World   \t\n\n\n\nSecond  line \r\n"
    cleaned = preprocess.clean_text(dirty)
    assert "\x00" not in cleaned
    assert "\n\n\n" not in cleaned
    for line in cleaned.split("\n"):
        assert "  " not in line


def test_convert_pdf_produces_nonempty_txt(sample_pdf: Path, tmp_path: Path):
    output_dir = tmp_path / "out"
    dest = preprocess.convert_file(sample_pdf, output_dir)
    assert dest is not None
    assert dest.exists()
    content = dest.read_text(encoding="utf-8")
    assert content.strip() != ""
    assert "GraphRAG" in content


def test_convert_docx_produces_nonempty_txt(sample_docx: Path, tmp_path: Path):
    output_dir = tmp_path / "out"
    dest = preprocess.convert_file(sample_docx, output_dir)
    assert dest is not None
    content = dest.read_text(encoding="utf-8")
    assert "GraphRAG" in content


def test_convert_txt_passthrough(tmp_path: Path):
    src = tmp_path / "sample.txt"
    src.write_text("단순 텍스트 파일입니다.", encoding="utf-8")
    output_dir = tmp_path / "out"
    dest = preprocess.convert_file(src, output_dir)
    assert dest is not None
    content = dest.read_text(encoding="utf-8")
    assert "단순 텍스트" in content


def test_hwp_is_skipped_with_warning(tmp_path: Path, caplog):
    src = tmp_path / "sample.hwp"
    src.write_bytes(b"dummy")
    output_dir = tmp_path / "out"
    with caplog.at_level("WARNING"):
        dest = preprocess.convert_file(src, output_dir)
    assert dest is None
    assert "HWP" in caplog.text


def test_convert_directory_handles_multiple_files(tmp_path: Path, sample_pdf: Path, sample_docx: Path):
    input_dir = tmp_path / "in"
    input_dir.mkdir()
    (input_dir / "a.pdf").write_bytes(sample_pdf.read_bytes())
    (input_dir / "b.docx").write_bytes(sample_docx.read_bytes())
    (input_dir / "c.txt").write_text("텍스트", encoding="utf-8")
    output_dir = tmp_path / "out"
    results = preprocess.convert_directory(input_dir, output_dir)
    assert len(results) == 3
    names = {r.name for r in results}
    assert names == {"a.txt", "b.txt", "c.txt"}


def test_convert_directory_missing_input_raises(tmp_path: Path):
    with pytest.raises(preprocess.PreprocessError):
        preprocess.convert_directory(tmp_path / "does_not_exist", tmp_path / "out")


def test_convert_directory_disambiguates_same_stem(tmp_path: Path, sample_pdf: Path, sample_docx: Path):
    input_dir = tmp_path / "in"
    input_dir.mkdir()
    (input_dir / "report.pdf").write_bytes(sample_pdf.read_bytes())
    (input_dir / "report.docx").write_bytes(sample_docx.read_bytes())
    output_dir = tmp_path / "out"
    results = preprocess.convert_directory(input_dir, output_dir)
    assert len(results) == 2
    names = {r.name for r in results}
    # sorted(iterdir())순으로 처리되므로 "report.docx"가 먼저 report.txt를 선점하고
    # 뒤이은 "report.pdf"가 충돌해 이름이 바뀐다.
    assert names == {"report.txt", "report__pdf.txt"}
    for r in results:
        assert r.read_text(encoding="utf-8").strip() != ""


def test_convert_directory_warns_on_unsupported_extension(tmp_path: Path, caplog):
    input_dir = tmp_path / "in"
    input_dir.mkdir()
    (input_dir / "data.xlsx").write_bytes(b"dummy")
    output_dir = tmp_path / "out"
    with caplog.at_level("WARNING"):
        results = preprocess.convert_directory(input_dir, output_dir)
    assert results == []
    assert "xlsx" in caplog.text.lower() or "지원하지" in caplog.text


def test_extract_docx_text_includes_table_content(tmp_path: Path):
    import docx

    document = docx.Document()
    document.add_paragraph("표 앞 문단입니다.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "이름"
    table.rows[0].cells[1].text = "박도현"
    dest = tmp_path / "with_table.docx"
    document.save(str(dest))

    text = preprocess.extract_docx_text(dest)
    assert "표 앞 문단입니다" in text
    assert "박도현" in text
