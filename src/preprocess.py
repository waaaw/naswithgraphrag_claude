#!/usr/bin/env python3
"""문서 전처리 (T04): pdf/docx/txt -> 정제된 UTF-8 txt.

파일당 1개의 .txt로 변환하여 --output(기본 ragproj/input/)에 저장한다.
hwp는 지원하지 않으며 경고 후 건너뛴다.
"""
from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path

logger = logging.getLogger("preprocess")

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
UNSUPPORTED_WARN_EXTENSIONS = {".hwp"}

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_MULTI_SPACE_TAB_RE = re.compile(r"[ \t]+")
_MULTI_BLANK_LINES_RE = re.compile(r"\n{3,}")


class PreprocessError(RuntimeError):
    """전처리 중 발생한, 사용자가 원인을 알 수 있는 명확한 오류."""


def clean_text(text: str) -> str:
    """제어문자·중복 공백·과도한 빈 줄을 정리한다."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _CONTROL_CHARS_RE.sub("", text)
    text = _MULTI_SPACE_TAB_RE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = _MULTI_BLANK_LINES_RE.sub("\n\n", text)
    return text.strip()


def extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def extract_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_txt_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def convert_file(src: Path, output_dir: Path, dest_name: str | None = None) -> Path | None:
    """단일 파일을 정제 UTF-8 txt로 변환한다. hwp 등 미지원 파일은 None 반환."""
    ext = src.suffix.lower()
    if ext == ".pdf":
        raw = extract_pdf_text(src)
    elif ext == ".docx":
        raw = extract_docx_text(src)
    elif ext == ".txt":
        raw = extract_txt_text(src)
    elif ext in UNSUPPORTED_WARN_EXTENSIONS:
        logger.warning("HWP는 지원하지 않습니다 (건너뜀): %s", src)
        return None
    else:
        raise PreprocessError(f"지원하지 않는 확장자입니다: {src}")

    cleaned = clean_text(raw)
    if not cleaned:
        logger.warning("추출된 텍스트가 비어 있습니다 (스캔 PDF 등일 수 있음): %s", src)
    output_dir.mkdir(parents=True, exist_ok=True)
    dest = output_dir / (dest_name or f"{src.stem}.txt")
    dest.write_text(cleaned, encoding="utf-8")
    logger.info("변환 완료: %s -> %s (%d자)", src, dest, len(cleaned))
    return dest


def convert_directory(input_dir: Path, output_dir: Path) -> list[Path]:
    if not input_dir.exists():
        raise PreprocessError(f"입력 경로가 없습니다: {input_dir}")
    results: list[Path] = []
    used_names: set[str] = set()
    for src in sorted(input_dir.iterdir()):
        if not src.is_file():
            continue
        ext = src.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS and ext not in UNSUPPORTED_WARN_EXTENSIONS:
            logger.warning("지원하지 않는 확장자라 건너뜁니다: %s", src)
            continue
        dest_name = f"{src.stem}.txt"
        if dest_name in used_names:
            # 같은 이름, 다른 확장자(예: report.pdf + report.docx)가 같은 디렉터리에
            # 있으면 출력 파일명이 충돌해 하나가 덮어써지므로 원본 확장자를 덧붙인다.
            disambiguated = f"{src.stem}__{ext.lstrip('.')}.txt"
            logger.warning(
                "출력 파일명이 중복되어 이름을 바꿉니다: %s -> %s", dest_name, disambiguated
            )
            dest_name = disambiguated
        dest = convert_file(src, output_dir, dest_name=dest_name)
        if dest is not None:
            used_names.add(dest_name)
            results.append(dest)
    return results


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
    parser = argparse.ArgumentParser(description="pdf/docx/txt -> 정제 UTF-8 txt 변환")
    parser.add_argument("--input", type=Path, required=True, help="원본 문서 디렉터리 또는 단일 파일")
    parser.add_argument(
        "--output", type=Path, default=Path("ragproj/input"),
        help="변환된 txt 저장 경로 (기본 ragproj/input)",
    )
    args = parser.parse_args()

    try:
        if args.input.is_dir():
            results = convert_directory(args.input, args.output)
            logger.info("총 %d개 파일 변환 완료", len(results))
        elif args.input.is_file():
            dest = convert_file(args.input, args.output)
            if dest is None:
                sys.exit(1)
        else:
            raise PreprocessError(f"입력 경로가 없습니다: {args.input}")
    except PreprocessError as e:
        logger.error(str(e))
        sys.exit(1)


if __name__ == "__main__":
    main()
